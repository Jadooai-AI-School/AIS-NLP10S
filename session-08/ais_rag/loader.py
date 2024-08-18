import os
import logging
from typing import List
import pymupdf  # PyMuPDF
from docx import Document as DocxDocument

# Configuring logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class Document:
    def __init__(self, page_content: str, metadata: dict):
        # Initializing Document with page content and metadata
        self.page_content = page_content
        self.metadata = metadata

    def __repr__(self):
        # Representing Document with truncated page content and metadata
        return f"Document(metadata={self.metadata}, page_content={self.page_content[:100]}...)"

def extract_metadata_from_pdf(file_path: str) -> dict:
    """Extracting metadata from a PDF file.

    Args:
        file_path: The path to the PDF file.

    Returns:
        A dictionary containing metadata extracted from the PDF.
    """
    try:
        pdf_document = pymupdf.open(file_path)  # Opening the PDF file
        metadata = pdf_document.metadata  # Extracting metadata
        return {
            "title": metadata.get('title', "").strip(),
            #"author": metadata.get('author', "").strip(),
            #"creation_date": metadata.get('creationDate', "").strip(),
        }
    except Exception as e:
        logging.error(f"Error extracting metadata from PDF file '{file_path}': {str(e)}")
        raise RuntimeError(f"Error extracting metadata from {file_path}") from e

def load_pdf(file_path: str) -> list:
    """Loading a PDF file and converting each page into Document objects.

    Args:
        file_path: The path to the PDF file.

    Returns:
        A list of Document objects containing page content and metadata.
    """
    # Checking if the PDF file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"The file {file_path} does not exist.")

    try:
        logging.info(f"Loading PDF file: {file_path}")
        pdf_document = pymupdf.open(file_path)  # Opening the PDF file
        metadata = extract_metadata_from_pdf(file_path)  # Extracting metadata
        documents = []
        for page_num in range(len(pdf_document)):  # Iterating through each page
            page = pdf_document.load_page(page_num)  # Loading the page
            text = page.get_text()  # Extracting text from the page
            combined_metadata = {**metadata, 'source': file_path, 'page_number': page_num + 1}
            documents.append(Document(page_content=text, metadata=combined_metadata))  # Creating Document object
        
        logging.info(f"Loaded PDF file: {file_path}")
        return documents
    except Exception as e:
        logging.error(f"Error loading PDF file '{file_path}': {str(e)}")
        raise RuntimeError(f"Error loading {file_path}") from e

def load_docx(file_path: str) -> List[Document]:
    """Loading a DOCX file and converting its content into a Document object.

    Args:
        file_path: The path to the DOCX file.

    Returns:
        A list containing a single Document object with the content and metadata.
    """
    logging.info(f"Loading DOCX file: {file_path}")
    doc = DocxDocument(file_path)  # Opening the DOCX file
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])  # Extracting text from paragraphs
    metadata = {'source': file_path}
    documents = [Document(page_content=text, metadata=metadata)]  # Creating Document object
    logging.info(f"Loaded DOCX file: {file_path}")
    return documents

def load_text(file_path: str) -> List[Document]:
    """Loading a text file and converting its content into a Document object.

    Args:
        file_path: The path to the text file.

    Returns:
        A list containing a single Document object with the content and metadata.
    """
    logging.info(f"Loading text file: {file_path}")
    # Checking if the text file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"The file {file_path} does not exist.")

    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()  # Reading the content of the file

    metadata = {'source': file_path}
    documents = [Document(page_content=text, metadata=metadata)]  # Creating Document object
    logging.info(f"Loaded text file: {file_path}")
    return documents

def load_documents(file_path: str) -> List[Document]:
    """Loading documents from a directory or a single file.

    Args:
        file_path: The path to the file or directory.

    Returns:
        A list of Document objects.
    """
    documents = []
    if os.path.isdir(file_path):  # Checking if the path is a directory
        logging.info(f"Loading all files in directory: {file_path}")
        for filename in os.listdir(file_path):  # Iterating through files in the directory
            full_path = os.path.join(file_path, filename)
            if os.path.isfile(full_path):  # Checking if the item is a file
                file_extension = filename.split(".")[-1].lower()
                if file_extension == "txt":
                    documents.extend(load_text(full_path))  # Loading text files
                elif file_extension == "pdf":
                    documents.extend(load_pdf(full_path))  # Loading PDF files
                elif file_extension == "docx":
                    documents.extend(load_docx(full_path))  # Loading DOCX files
                else:
                    logging.warning(f"Unsupported file extension: {file_extension}")
    elif os.path.isfile(file_path):  # Checking if the path is a single file
        logging.info(f"Loading single file: {file_path}")
        file_extension = file_path.split(".")[-1].lower()
        if file_extension == "txt":
            documents = load_text(file_path)  # Loading text file
        elif file_extension == "pdf":
            documents = load_pdf(file_path)  # Loading PDF file
        elif file_extension == "docx":
            documents = load_docx(file_path)  # Loading DOCX file
        else:
            raise ValueError(f"Unsupported file extension: {file_extension}")
    else:
        raise ValueError(f"The path {file_path} is neither a file nor a directory.")
    
    logging.info(f"Loaded documents: {len(documents)}")
    return documents

if __name__ == "__main__":
    path = "data"  # Can be a directory or a single file
    os.system('cls' if os.name == 'nt' else 'clear')  # Clearing the console
    documents = load_documents(path)  # Loading documents from the specified path
    for doc in documents:
        print(doc.metadata)  # Printing document metadata
        print('-----------------')
        print(doc.page_content)  # Printing document content
